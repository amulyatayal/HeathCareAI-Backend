#!/usr/bin/env python3
"""Generate Nutrition_Meal_Plan_Pipeline_Flow.docx using editable Word table blocks.

Each "block" in a flowchart is a single-cell table with shaded background.
Arrows are plain paragraphs containing arrow glyphs.
Branches use multi-column tables where each column holds an arrow + block.
Everything is native Word — no images — so the user can edit any block in place.
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
OUTPUT = DOCS_DIR / "Nutrition_Meal_Plan_Pipeline_Flow.docx"

# Block fill colors (hex, no #)
COLOR_INPUT = "DBEAFE"      # blue   — inputs / CLI
COLOR_PROCESS = "D1FAE5"    # green  — agents / processing
COLOR_DECISION = "FEF3C7"   # amber  — decisions
COLOR_DOC = "EDE9FE"        # violet — reference docs
COLOR_OUTPUT = "FFE4E6"     # rose   — outputs
COLOR_BORDER = "334155"


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_borders(cell, hex_color: str = COLOR_BORDER, sz: int = 8) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), hex_color)
        borders.append(b)
    tc_pr.append(borders)


def _set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = OxmlElement("w:tcW")
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")
    tc_pr.append(tc_w)


def _fill_block_cell(
    cell,
    title: str,
    body: str = "",
    color: str = COLOR_PROCESS,
    title_size: int = 11,
    body_size: int = 9,
) -> None:
    cell.text = ""
    _set_cell_shading(cell, color)
    _set_cell_borders(cell, COLOR_BORDER, sz=10)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_title = cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(0)
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(title_size)
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    if body:
        p_body = cell.add_paragraph()
        p_body.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_body.paragraph_format.space_before = Pt(2)
        p_body.paragraph_format.space_after = Pt(2)
        run = p_body.add_run(body)
        run.font.size = Pt(body_size)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)


def add_block(
    doc: Document,
    title: str,
    body: str = "",
    color: str = COLOR_PROCESS,
    width_in: float = 4.6,
) -> None:
    """Single editable block centered on the page."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.rows[0].cells[0]
    _set_cell_width(cell, width_in)
    _fill_block_cell(cell, title, body, color)


def add_arrow(doc: Document, label: str = "", glyph: str = "↓") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(glyph)
    run.font.size = Pt(18)
    run.bold = True
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    if label:
        run2 = p.add_run(f"  {label}")
        run2.italic = True
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def add_branch_row(
    doc: Document,
    branches: List[dict],
    col_width_in: float = 2.2,
) -> None:
    """Horizontal row of blocks (one per branch).

    Each branch dict: {"title": str, "body": str, "color": str, "label": Optional[str]}.
    The optional 'label' is shown above the block (e.g. "Yes" / "No").
    """
    n = len(branches)
    table = doc.add_table(rows=2, cols=n)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Top row = arrow + label
    for i, br in enumerate(branches):
        cell = table.rows[0].cells[i]
        _set_cell_width(cell, col_width_in)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("↓")
        run.font.size = Pt(16)
        run.bold = True
        run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        if br.get("label"):
            run2 = p.add_run(f"  {br['label']}")
            run2.italic = True
            run2.font.size = Pt(9)

    # Bottom row = block content
    for i, br in enumerate(branches):
        cell = table.rows[1].cells[i]
        _set_cell_width(cell, col_width_in)
        _fill_block_cell(
            cell,
            br.get("title", ""),
            br.get("body", ""),
            br.get("color", COLOR_PROCESS),
            title_size=10,
            body_size=8,
        )


def add_decision(doc: Document, title: str, body: str = "") -> None:
    """Diamond-style decision (we render as an amber block with a leading '◆')."""
    add_block(doc, f"◆  {title}", body, color=COLOR_DECISION, width_in=4.6)


# ---------------------------------------------------------------------------
# Plain helpers
# ---------------------------------------------------------------------------
def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True


def add_bullets(doc: Document, items: List[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
        _set_cell_shading(hdr_cells[i], "F1F5F9")
    for ri, row in enumerate(rows):
        row_cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            row_cells[ci].text = str(val)
    doc.add_paragraph()


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)


def add_legend(doc: Document) -> None:
    add_heading(doc, "Block legend", 2)
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    legend = [
        ("Input", COLOR_INPUT),
        ("Process / Agent", COLOR_PROCESS),
        ("Decision", COLOR_DECISION),
        ("Reference doc", COLOR_DOC),
        ("Output", COLOR_OUTPUT),
    ]
    for i, (label, color) in enumerate(legend):
        cell = table.rows[0].cells[i]
        _set_cell_width(cell, 1.3)
        _fill_block_cell(cell, label, "", color, title_size=10)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Main flowchart sections
# ---------------------------------------------------------------------------
def section_main_pipeline(doc: Document) -> None:
    add_heading(doc, "1. Main pipeline", 1)
    add_para(
        doc,
        "Linear flow from CLI invocation through to the saved meal-plan markdown. "
        "Click any block to edit its text directly.",
    )
    add_caption(doc, "Each colored cell below is a Word table cell — edit text in place.")

    blocks = [
        ("1. CLI / Test Driver",
         "tests/nutrition_meal_plan.py\n--patient · --espen · --question",
         COLOR_INPUT),
        ("2. build_nutrition_context",
         "Loads patient JSON, ESPEN KB, summaries\n→ PipelineContext.metadata",
         COLOR_PROCESS),
        ("3. StageAgentV2",
         "Merge user_data · BMI · clinical categories\nLLM journey-stage classification",
         COLOR_PROCESS),
        ("4. RetrievalAgent",
         "OpenSearch: nutrition_assistant + medical_all_kb\n(optional — strict_rag = False)",
         COLOR_PROCESS),
        ("5. ReasoningAgent (nutrition)",
         "ESPEN-Energy-KnowledgeBase.json + Bedrock LLM\nDerives all numeric targets + meal plan",
         COLOR_PROCESS),
        ("6. Output",
         "render_result() → tests/meal_plans/meal_plan_<id>_<ts>.md",
         COLOR_OUTPUT),
    ]
    for i, (title, body, color) in enumerate(blocks):
        if i > 0:
            add_arrow(doc)
        add_block(doc, title, body, color)

    doc.add_paragraph()
    add_para(doc, "Side branch from step 3 if mandatory user_data is missing:", italic=True)
    add_branch_row(
        doc,
        [
            {"title": "◆ Mandatory data?",
             "body": "Required fields present?",
             "color": COLOR_DECISION,
             "label": "check"},
            {"title": "Abort pipeline",
             "body": "No meal plan produced",
             "color": COLOR_OUTPUT,
             "label": "No"},
        ],
        col_width_in=2.6,
    )

    doc.add_paragraph()
    add_bullets(
        doc,
        [
            "Step 3 abort path is skipped in the test driver via skip_user_data_confirmations=True.",
            "Step 4 returns 0 chunks when OpenSearch is unavailable; reasoning still runs.",
            "Nutrition ReasoningAgent timeout: 180 s · max_tokens: 3000.",
        ],
    )


def section_inputs(doc: Document) -> None:
    add_heading(doc, "2. Inputs and source files", 1)
    add_para(
        doc,
        "Three input streams flow into build_nutrition_context, which packs them "
        "into PipelineContext.metadata for every downstream agent.",
    )

    add_caption(doc, "Inputs (left) — Aggregator (center) — Reference docs (right)")
    add_branch_row(
        doc,
        [
            {"title": "Patient JSON", "body": "24 fields\nPatient_JSON_Records/*.json", "color": COLOR_INPUT},
            {"title": "--question", "body": "Meal-plan request text", "color": COLOR_INPUT},
            {"title": "--espen dir", "body": "data/EspenGuideline/", "color": COLOR_INPUT},
        ],
        col_width_in=2.0,
    )
    add_arrow(doc, label="loaded by build_nutrition_context")
    add_block(
        doc,
        "build_nutrition_context()",
        "Returns PipelineContext with metadata:\n"
        "user_data · patient_data · espen_nutrition_kb · espen_guidelines",
        COLOR_PROCESS,
        width_in=5.6,
    )
    add_arrow(doc, label="reference docs attached")
    add_branch_row(
        doc,
        [
            {"title": "ESPEN-Energy-\nKnowledgeBase.json",
             "body": "PRIMARY rules:\nkcal · protein · macros · safety",
             "color": COLOR_DOC},
            {"title": "Espen-Guideline.json\nEspen-Summary.json",
             "body": "Supplementary background",
             "color": COLOR_DOC},
            {"title": "clinical_thresholds.json",
             "body": "BMI · waist · grip · PG-SGA labels",
             "color": COLOR_DOC},
            {"title": "stage_hierarchy.json",
             "body": "Journey stage taxonomy",
             "color": COLOR_DOC},
        ],
        col_width_in=1.55,
    )

    doc.add_paragraph()
    add_heading(doc, "2.1 Patient JSON fields (24 keys)", 2)
    add_para(
        doc,
        "Clinical inputs only — no pre-computed kcal/protein. The LLM derives targets from ESPEN rules.",
    )
    add_table(
        doc,
        ["Field group", "Examples", "Used for"],
        [
            ["Body metrics", "Height_cm, Weight_kg, BMI, Waist, Hand_Grip", "kcal/kg math + categorization"],
            ["Cancer / treatment", "Cancer_Stage, Treatment_Modality, Chemotherapy_Regimen", "Subgroup overrides (ESPEN Cx)"],
            ["Nutrition status", "PG_SGA_Score, Nutrition_Impact_Symptoms", "PG-SGA label + meal adjustments"],
            ["Demographics / diet", "Region_India, Dietary_Pattern", "Cultural meal personalization"],
        ],
    )


def section_stage_agent(doc: Document) -> None:
    add_heading(doc, "3. StageAgentV2 — categorization (no kcal)", 1)
    add_para(
        doc,
        "StageAgent merges every available source of user_data, computes BMI if missing, "
        "applies clinical_thresholds.json to produce category labels, and classifies the "
        "patient's journey stage with an LLM call. It does NOT compute energy or protein targets.",
    )
    add_caption(doc, "Sources merge into user_data, then split into category + stage outputs.")

    add_branch_row(
        doc,
        [
            {"title": "Patient JSON", "body": "Loaded record", "color": COLOR_INPUT},
            {"title": "DynamoDB biomarkers", "body": "Optional", "color": COLOR_INPUT},
            {"title": "Chat extraction", "body": "Optional", "color": COLOR_INPUT},
            {"title": "Caller user_data", "body": "Optional", "color": COLOR_INPUT},
        ],
        col_width_in=1.55,
    )
    add_arrow(doc, label="merged (last source wins)")
    add_block(
        doc,
        "merged_user_data",
        "+ compute BMI if missing",
        COLOR_PROCESS,
        width_in=4.6,
    )
    add_arrow(doc)
    add_branch_row(
        doc,
        [
            {"title": "categorize_patient()",
             "body": "Reads clinical_thresholds.json\n→ clinical_categories",
             "color": COLOR_PROCESS},
            {"title": "LLM stage classifier",
             "body": "Reads stage_hierarchy.json\n→ stage_result",
             "color": COLOR_PROCESS},
        ],
        col_width_in=2.8,
    )
    add_arrow(doc)
    add_branch_row(
        doc,
        [
            {"title": "clinical_categories",
             "body": "BMI, waist, grip, PG-SGA labels",
             "color": COLOR_OUTPUT},
            {"title": "stage_result",
             "body": "Active treatment · Survivorship · …",
             "color": COLOR_OUTPUT},
        ],
        col_width_in=2.8,
    )

    doc.add_paragraph()
    add_para(doc, "Mandatory-fields gate (skipped in test driver):", italic=True)
    add_branch_row(
        doc,
        [
            {"title": "◆ Mandatory data?",
             "body": "Required fields present?",
             "color": COLOR_DECISION,
             "label": ""},
            {"title": "Abort",
             "body": "Pipeline halts",
             "color": COLOR_OUTPUT,
             "label": "No"},
        ],
        col_width_in=2.6,
    )

    doc.add_paragraph()
    add_table(
        doc,
        ["Metric", "Config file", "Output label examples"],
        [
            ["BMI", "clinical_thresholds.json (indian_icmr)", "Obese Class I, Underweight"],
            ["Waist", "clinical_thresholds.json + sex", "Central Obesity"],
            ["Hand grip", "clinical_thresholds.json + sex + age", "Low, Normal"],
            ["PG-SGA", "clinical_thresholds.json", "B - Moderately / C - Severely Malnourished"],
        ],
    )


def section_reasoning(doc: Document) -> None:
    add_heading(doc, "4. ReasoningAgent — final nutrition values", 1)
    add_para(
        doc,
        "All numeric nutrition targets (kcal/day, protein g/day, macro split, vitamins) "
        "are computed by the Bedrock LLM following ESPEN-Energy-KnowledgeBase.json. "
        "Every other input is supplementary context.",
    )

    add_caption(doc, "Authoritative source (top) and supplementary context (left) feed the agent.")
    add_block(
        doc,
        "AUTHORITATIVE — ESPEN-Energy-KnowledgeBase.json",
        "Energy [B2-1] · Protein [B2-2] · Macros [B2-3] · Vitamins [B2-4]\n"
        "Refeeding [B3-4] · Subgroups [C1-C6] · Pseudocode decision tree",
        COLOR_DOC,
        width_in=5.8,
    )
    add_arrow(doc, label="primary rules")
    add_block(
        doc,
        "ReasoningAgent (Bedrock LLM)",
        "Walks the rules in order:\nrefeeding → energy → protein → macros → vitamins → meals",
        COLOR_PROCESS,
        width_in=5.8,
    )
    add_arrow(doc, label="produces")
    add_branch_row(
        doc,
        [
            {"title": "Nutrition Target Derivation",
             "body": "Rules + arithmetic shown",
             "color": COLOR_OUTPUT},
            {"title": "Daily Nutrition Targets",
             "body": "kcal · protein · micros",
             "color": COLOR_OUTPUT},
            {"title": "Meal Plan",
             "body": "Breakfast … Dinner",
             "color": COLOR_OUTPUT},
        ],
        col_width_in=2.0,
    )

    doc.add_paragraph()
    add_para(doc, "Supplementary context (also passed in metadata):", italic=True)
    add_branch_row(
        doc,
        [
            {"title": "patient_data block",
             "body": "Treatment, symptoms, diet",
             "color": COLOR_INPUT},
            {"title": "clinical_categories",
             "body": "BMI, PG-SGA labels (verbatim)",
             "color": COLOR_PROCESS},
            {"title": "espen_guidelines",
             "body": "Summary JSON files",
             "color": COLOR_DOC},
            {"title": "OpenSearch chunks",
             "body": "[Source N] (optional)",
             "color": COLOR_INPUT},
            {"title": "stage + question",
             "body": "Journey context",
             "color": COLOR_INPUT},
        ],
        col_width_in=1.25,
    )

    doc.add_paragraph()
    add_heading(doc, "4.1 Document authority", 2)
    add_table(
        doc,
        ["Document", "Path", "Role"],
        [
            ["ESPEN Energy KB (PRIMARY)", "data/EspenGuideline/ESPEN-Energy-KnowledgeBase.json", "All numeric targets + citations"],
            ["ESPEN summaries", "Espen-Guideline.json, Espen-Summary.json", "Background only"],
            ["Clinical thresholds", "config/clinical_thresholds.json", "Category labels only"],
            ["Stage hierarchy", "data/stage_hierarchy.json", "Journey stage"],
            ["Patient JSON", "Patient_JSON_Records/*.json", "Measurements + treatment"],
            ["OpenSearch", "nutrition_assistant, medical_all_kb", "Optional [Source N]"],
        ],
    )


def section_espen_decision(doc: Document) -> None:
    add_heading(doc, "5. ESPEN energy target — decision flow", 1)
    add_para(
        doc,
        "Simplified branch logic that the LLM follows when computing daily kcal. "
        "Edit any block below to reflect updated rules.",
    )

    add_block(doc, "Start", "Patient weight · BMI · treatment phase", COLOR_INPUT, width_in=4.6)
    add_arrow(doc)

    add_decision(doc, "Refeeding risk?  [B3-4]",
                 "BMI < 18.5 or > 10–15% weight loss or low intake")
    add_branch_row(
        doc,
        [
            {"title": "5–10 kcal/kg",
             "body": "Ramp over 4–7 days",
             "color": COLOR_PROCESS,
             "label": "Yes"},
            {"title": "Continue",
             "body": "No refeed gate",
             "color": COLOR_PROCESS,
             "label": "No"},
        ],
        col_width_in=2.6,
    )
    add_arrow(doc)

    add_decision(doc, "BMI > 30?  [B2-1]", "Obesity adjust")
    add_branch_row(
        doc,
        [
            {"title": "22–25 kcal/kg",
             "body": "× adjusted body weight",
             "color": COLOR_PROCESS,
             "label": "Yes"},
            {"title": "25–30 kcal/kg",
             "body": "× actual body weight",
             "color": COLOR_PROCESS,
             "label": "No"},
        ],
        col_width_in=2.6,
    )
    add_arrow(doc)

    add_decision(doc, "Cachexia / high CRP?  [B2-3]",
                 "Apply upper end of range + fat-shifted macros")
    add_arrow(doc)

    add_block(
        doc,
        "Apply subgroup overrides  [Cx]",
        "C3 surgery · C6 chemo · C1 head-and-neck · C2 thoracic · C4 GI · C5 hematologic",
        COLOR_DOC,
        width_in=5.6,
    )
    add_arrow(doc)
    add_block(
        doc,
        "Final kcal/day",
        "weight × selected kcal/kg",
        COLOR_OUTPUT,
        width_in=4.6,
    )

    doc.add_paragraph()
    add_table(
        doc,
        ["Step", "ESPEN rule", "Typical formula"],
        [
            ["Refeeding", "B3-4", "5–10 kcal/kg → ramp over 4–7 days"],
            ["Obesity", "B2-1", "22–25 kcal/kg × adjusted body weight if BMI > 30"],
            ["Default", "B2-1", "25–30 kcal/kg × actual body weight"],
            ["Cachexia modifier", "B2-3", "Upper end of range; fat-shifted macros"],
            ["Protein", "B2-2", "1.0–1.5 g/kg (up to ~2.0 g/kg if appropriate)"],
        ],
    )


def section_cli_summary(doc: Document) -> None:
    add_heading(doc, "6. CLI reference", 1)
    add_table(
        doc,
        ["Argument", "Required?", "Default"],
        [
            ["--patient", "No", "PATIENT_DATA_PATH (IN-BC-4001.json)"],
            ["--espen", "No", "data/EspenGuideline/"],
            ["--question", "No", "DEFAULT_QUESTION"],
            ["--save / --no-save", "No", "Auto-save to tests/meal_plans/"],
        ],
    )

    add_heading(doc, "7. Summary", 1)
    add_para(
        doc,
        "Patient JSON and question supply raw clinical data. StageAgent produces "
        "authoritative category labels from clinical_thresholds.json. "
        "ReasoningAgent computes all nutrition targets and meals by applying "
        "ESPEN-Energy-KnowledgeBase.json via Bedrock LLM — not from pre-stored "
        "values in the patient file.",
        bold=True,
    )
    add_para(doc, "Regenerate: python docs/generate_nutrition_flow_doc.py")


# ---------------------------------------------------------------------------
# Document
# ---------------------------------------------------------------------------
def build_document() -> Document:
    doc = Document()
    title = doc.add_heading("Nutrition Meal Plan Pipeline — Editable Block Flow", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_para(
        doc,
        "Every flowchart below is built from native Word tables. Each colored cell "
        "is a single block: click inside any cell to edit text; right-click → Table "
        "Properties to change size or shading.",
    )
    add_legend(doc)

    section_main_pipeline(doc)
    doc.add_page_break()
    section_inputs(doc)
    doc.add_page_break()
    section_stage_agent(doc)
    doc.add_page_break()
    section_reasoning(doc)
    doc.add_page_break()
    section_espen_decision(doc)
    doc.add_page_break()
    section_cli_summary(doc)

    return doc


def main() -> None:
    print("Building Word document with editable block diagrams...")
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
